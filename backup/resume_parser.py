import os
import re
import json
import datetime
import pdfplumber
import spacy
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

nlp = spacy.load("en_core_web_lg")

# ── Section heading keywords ──────────────────────────────────────────────────
SECTION_HEADERS = {
    "experience":       r"(work\s*experience|professional\s*experience|employment|experience)",
    "education":        r"(education|academic|qualification)",
    "graduation":       r"(graduation|under\s*graduate|bachelor|b\.?tech|b\.?e\.?|b\.?sc)",
    "post_graduation":  r"(post\s*graduation|post\s*graduate|master|m\.?tech|m\.?e\.?|m\.?sc|mba|mca)",
    "skills":           r"(skills|technical\s*skills|core\s*competencies|key\s*skills)",
    "projects":         r"(projects?|personal\s*projects?)",
    "certifications":   r"(certifications?|courses?|training)",
    "languages":        r"(languages?|language\s*known|spoken\s*languages?)",
    "social":           r"(social|linkedin|github|portfolio|online\s*profile)",
    "summary":          r"(summary|objective|profile|about\s*me)",
}

DEGREE_PATTERNS = [
    r"\b(B\.?Tech|B\.?E\.?|Bachelor\s+of\s+\w+|M\.?Tech|M\.?E\.?|Master\s+of\s+\w+|MBA|MCA|BCA|B\.?Sc\.?|M\.?Sc\.?|Ph\.?D\.?|Diploma)\b",
]

DESIGNATION_PATTERNS = [
    r"\b(Software\s+Engineer|Senior\s+Engineer|Junior\s+Engineer|Full\s*Stack\s+Developer|"
    r"Front[\s-]*end\s+Developer|Back[\s-]*end\s+Developer|Data\s+Scientist|"
    r"Data\s+Analyst|DevOps\s+Engineer|QA\s+Engineer|Product\s+Manager|"
    r"Project\s+Manager|Team\s+Lead|Tech\s+Lead|Assistant\s+Manager|"
    r"Operations\s+Manager|Business\s+Analyst|System\s+Analyst|"
    r"Web\s+Developer|Mobile\s+Developer|Android\s+Developer|iOS\s+Developer|"
    r"Machine\s+Learning\s+Engineer|AI\s+Engineer|Cloud\s+Engineer|"
    r"Intern|Trainee|Consultant|Architect|Founder|Co-Founder|Director|"
    r"Manager|Developer|Programmer|Analyst|Designer|Mentor|Faculty)\b",
]

SKILLS_DB = [
    "Python","Java","C","C++","C#","JavaScript","TypeScript","PHP","Ruby","Go","Rust","Kotlin","Swift",
    "HTML","HTML5","CSS","CSS3","SASS","SCSS","Bootstrap","Tailwind",
    "React","ReactJS","Angular","Vue","Node.js","NodeJS","Express","Django","Flask","FastAPI","Spring","Spring Boot",
    "MySQL","PostgreSQL","MongoDB","SQLite","Oracle","Redis","Firebase","DynamoDB","SQL","NoSQL",
    "AWS","Azure","GCP","Docker","Kubernetes","Jenkins","CI/CD","Git","GitHub","GitLab","Bitbucket",
    "REST","RESTful","GraphQL","API","Microservices","SOAP",
    "Machine Learning","Deep Learning","NLP","TensorFlow","PyTorch","Keras","Scikit-learn","Pandas","NumPy","Matplotlib",
    "Android","iOS","React Native","Flutter","Jetpack Compose",
    "Linux","Windows","Unix","Bash","Shell","PowerShell",
    "Agile","Scrum","Jira","Confluence","Trello",
    "Selenium","JUnit","Pytest","Postman","Swagger",
    "Excel","Word","PowerPoint","Tableau","Power BI",
    "Figma","Adobe XD","Photoshop","Illustrator",
    "Wordpress","Shopify","Magento","Drupal",
    "Hibernate","JPA","JDBC","Maven","Gradle","Webpack","Babel",
    "XML","JSON","YAML","Markdown",
    "Algorithms","Data Structures","OOP","OOPS","Design Patterns","SDLC",
    "Communication","Leadership","Teamwork","Problem Solving","Analytical",
]

DATE_PATTERN = r"(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s*[\.,]?\s*\d{4}"
YEAR_RANGE   = r"\d{4}\s*[-–—to]+\s*(\d{4}|present|current|now)"


# ── Text extraction ───────────────────────────────────────────────────────────
def extract_text_pdf(path):
    pages = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text(x_tolerance=2, y_tolerance=2)
            if t:
                pages.append(t)
    return "\n".join(pages)


def extract_text_docx(path):
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_text(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_text_pdf(path)
    elif ext in (".docx", ".doc"):
        return extract_text_docx(path)
    return ""


# ── Field extractors ──────────────────────────────────────────────────────────
def extract_email(text):
    m = re.search(r"[\w.+-]+@[\w.-]+\.\w{2,}", text)
    return m.group(0) if m else None


def extract_phone(text):
    for pat in [r"\+91[-\s]?\d{10}", r"\b\d{10}\b", r"\d{3}[-.\s]\d{3}[-.\s]\d{4}"]:
        m = re.search(pat, text)
        if m:
            return m.group(0).strip()
    return None


def extract_name(text, email=None):
    lines = [l.strip() for l in text.splitlines() if l.strip()]

    # Skip lines that look like headers/contact info
    skip = re.compile(r"@|\d{5,}|linkedin|github|http|curriculum|resume|cv|page\s*\d", re.I)

    candidates = []
    for line in lines[:10]:
        if skip.search(line):
            continue
        words = line.split()
        if 2 <= len(words) <= 5 and all(w[0].isupper() or w[0] == "/" for w in words if w):
            candidates.append(line)

    if candidates:
        return candidates[0]

    # Fallback: spaCy PERSON entity from first 500 chars
    doc = nlp(text[:500])
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            return ent.text.strip()

    return None


def extract_skills(text):
    found = set()
    text_lower = text.lower()
    for skill in SKILLS_DB:
        pattern = r"\b" + re.escape(skill.lower()) + r"\b"
        if re.search(pattern, text_lower):
            found.add(skill)
    return sorted(found)


def extract_degrees(text):
    found = []
    for pat in DEGREE_PATTERNS:
        matches = re.findall(pat, text, re.I)
        found.extend(matches)
    return list(dict.fromkeys(found)) or None  # deduplicate, preserve order


def extract_designations(text):
    found = []
    for pat in DESIGNATION_PATTERNS:
        matches = re.findall(pat, text, re.I)
        found.extend(matches)
    return list(dict.fromkeys(found)) or None


def extract_college(text):
    patterns = [
        r"((?:IIT|NIT|BITS|MIT|VIT|SRM|Amity|Anna\s+University|Mumbai\s+University|"
        r"Delhi\s+University|Pune\s+University|[\w\s]+(?:University|Institute|College|"
        r"School\s+of|Academy|Polytechnic)[\w\s,]*?))",
    ]
    for pat in patterns:
        m = re.search(pat, text, re.I)
        if m:
            name = m.group(1).strip().rstrip(",.")
            if len(name) > 5:
                return name
    return None


def split_sections(text):
    """Split resume text into named sections."""
    all_headers = "|".join(SECTION_HEADERS.values())
    header_re = re.compile(
        r"^[ \t]*(" + all_headers + r")[ \t]*[:\-]?[ \t]*$",
        re.I | re.M,
    )

    sections = {}
    matches = list(header_re.finditer(text))

    for i, m in enumerate(matches):
        label = m.group(0).strip().lower()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        content = text[start:end].strip()

        for key, pat in SECTION_HEADERS.items():
            if re.search(pat, label, re.I):
                sections.setdefault(key, content)
                break

    return sections


def extract_experience_grouped(exp_text):
    """Group experience into structured entries: company, role, duration, details."""
    if not exp_text:
        return None

    date_re = re.compile(
        r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s,]*\d{4}|\d{4})"
        r"\s*[-–—to]+\s*"
        r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s,]*\d{4}|\d{4}|present|current|now)",
        re.I,
    )
    desig_re = re.compile("|".join(DESIGNATION_PATTERNS), re.I)

    lines = [l.strip() for l in exp_text.splitlines() if l.strip()]
    entries, current = [], {}

    for line in lines:
        date_match = date_re.search(line)
        desig_match = desig_re.search(line)

        if date_match and (desig_match or current.get("role")):
            if current:
                entries.append(current)
            current = {
                "role": desig_match.group(0).strip() if desig_match else current.get("role", ""),
                "duration": f"{date_match.group(1).strip()} - {date_match.group(2).strip()}",
                "company": "",
                "details": [],
            }
        elif desig_match and not current.get("role"):
            current["role"] = desig_match.group(0).strip()
        elif current:
            if not current.get("company") and len(line.split()) <= 6 and line[0].isupper():
                current["company"] = line
            else:
                current.setdefault("details", []).append(line)

    if current:
        entries.append(current)

    # Fallback: return raw lines if no structured entries found
    if not entries:
        return [{"details": lines}]

    return entries


def extract_experience_entries(exp_text):
    if not exp_text:
        return None
    lines = [l.strip() for l in exp_text.splitlines() if l.strip()]
    return lines if lines else None


def calculate_experience_years(exp_text):
    if not exp_text:
        return 0
    total_months = 0
    ranges = re.findall(
        r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s,]*\d{4})"
        r"\s*[-–—to]+\s*"
        r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s,]*\d{4}|present|current|now)",
        exp_text, re.I,
    )

    month_map = {
        "jan":1,"feb":2,"mar":3,"apr":4,"may":5,"jun":6,
        "jul":7,"aug":8,"sep":9,"oct":10,"nov":11,"dec":12,
    }

    import datetime
    now = datetime.datetime.now()

    for start_str, end_str in ranges:
        try:
            sm = re.search(r"(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)", start_str, re.I)
            sy = re.search(r"\d{4}", start_str)
            if not sm or not sy:
                continue
            s_month = month_map[sm.group(0).lower()]
            s_year  = int(sy.group(0))

            if re.search(r"present|current|now", end_str, re.I):
                e_month, e_year = now.month, now.year
            else:
                em = re.search(r"(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)", end_str, re.I)
                ey = re.search(r"\d{4}", end_str)
                if not em or not ey:
                    continue
                e_month = month_map[em.group(0).lower()]
                e_year  = int(ey.group(0))

            total_months += (e_year - s_year) * 12 + (e_month - s_month)
        except Exception:
            continue

    return round(total_months / 12, 2) if total_months else 0


def extract_company_names(exp_text):
    if not exp_text:
        return None
    doc = nlp(exp_text[:2000])
    orgs = [ent.text.strip() for ent in doc.ents if ent.label_ == "ORG"]
    return list(dict.fromkeys(orgs)) or None


def extract_languages(text):
    known_langs = [
        "English", "Hindi", "Tamil", "Telugu", "Kannada", "Malayalam", "Marathi",
        "Bengali", "Gujarati", "Punjabi", "Urdu", "French", "German", "Spanish",
        "Japanese", "Chinese", "Arabic", "Portuguese", "Russian",
    ]
    found = []
    for lang in known_langs:
        if re.search(r"\b" + lang + r"\b", text, re.I):
            found.append(lang)
    return found or None


def extract_social_links(text):
    links = {}
    linkedin = re.search(r"(linkedin\.com/in/[\w\-]+|linkedin\.com/pub/[\w\-/]+)", text, re.I)
    github   = re.search(r"github\.com/[\w\-]+", text, re.I)
    portfolio = re.search(r"https?://(?!linkedin|github)[\w./-]+", text, re.I)
    if linkedin:  links["linkedin"]  = linkedin.group(0)
    if github:    links["github"]    = github.group(0)
    if portfolio: links["portfolio"] = portfolio.group(0)
    return links or None


def extract_graduation(edu_text):
    """Extract UG degree details."""
    ug_re = re.compile(
        r"(B\.?Tech|B\.?E\.?|B\.?Sc\.?|BCA|Bachelor[\w\s]+)",
        re.I,
    )
    m = ug_re.search(edu_text or "")
    return {"degree": m.group(0).strip()} if m else None


def extract_post_graduation(edu_text):
    """Extract PG degree details."""
    pg_re = re.compile(
        r"(M\.?Tech|M\.?E\.?|M\.?Sc\.?|MCA|MBA|Master[\w\s]+|Ph\.?D\.?)",
        re.I,
    )
    m = pg_re.search(edu_text or "")
    return {"degree": m.group(0).strip()} if m else None


def extract_certifications(cert_text):
    if not cert_text:
        return None
    lines = [l.strip() for l in cert_text.splitlines() if l.strip()]
    return lines or None


def extract_projects(proj_text):
    if not proj_text:
        return None
    lines = [l.strip() for l in proj_text.splitlines() if l.strip()]
    return lines or None


def count_pages(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        with pdfplumber.open(path) as pdf:
            return len(pdf.pages)
    return 1


# ── Main parser ───────────────────────────────────────────────────────────────
def parse_resume(path):
    text = extract_text(path)
    if not text.strip():
        return {}

    sections = split_sections(text)
    exp_text  = sections.get("experience", "")
    edu_text  = sections.get("education", text)
    email     = extract_email(text)

    return {
        "Profile_Personal_Details": {
            "name":          extract_name(text, email),
            "email":         email,
            "mobile_number": extract_phone(text),
            "social_links":  extract_social_links(text),
            "no_of_pages":   count_pages(path),
        },
        "Profile_Experience": {
            "designation":      extract_designations(exp_text or text),
            "company_names":    extract_company_names(exp_text),
            "total_experience": calculate_experience_years(exp_text),
            "experience_groups": extract_experience_grouped(exp_text),
        },
        "Profile_Graduation":      extract_graduation(edu_text),
        "Profile_Post_Graduation": extract_post_graduation(edu_text),
        "Profile_Certification":   extract_certifications(sections.get("certifications", "")),
        "Profile_Technical_Skills": extract_skills(text),
        "Profile_Projects":        extract_projects(sections.get("projects", "")),
        "Profile_Languages":       extract_languages(sections.get("languages", text)),
        "Profile_Social_Network":  extract_social_links(text),
    }


# ── Export helpers ────────────────────────────────────────────────────────────
def _flat_lines(data, indent=0):
    """Recursively flatten parsed data into (label, value) pairs for export."""
    lines = []
    prefix = "  " * indent
    if isinstance(data, dict):
        for k, v in data.items():
            if isinstance(v, (dict, list)):
                lines.append((prefix + str(k), ""))
                lines.extend(_flat_lines(v, indent + 1))
            else:
                lines.append((prefix + str(k), str(v) if v is not None else ""))
    elif isinstance(data, list):
        for item in data:
            if isinstance(item, dict):
                lines.extend(_flat_lines(item, indent))
                lines.append(("", ""))  # blank separator
            else:
                lines.append((prefix + "•", str(item)))
    return lines


def export_to_pdf(parsed, out_path):
    doc = SimpleDocTemplate(out_path, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("SectionTitle", parent=styles["Heading2"], textColor=colors.HexColor("#1a3c5e"), spaceAfter=4)
    normal = styles["Normal"]
    story = []

    for section_title, section_data in parsed.items():
        story.append(Paragraph(section_title.replace("_", " "), title_style))
        rows = _flat_lines(section_data)
        for label, value in rows:
            if label == "" and value == "":
                story.append(Spacer(1, 4))
            elif value:
                story.append(Paragraph(f"<b>{label.strip()}:</b> {value}", normal))
            elif label:
                story.append(Paragraph(f"<b>{label.strip()}</b>", normal))
        story.append(Spacer(1, 10))

    doc.build(story)


def export_to_word(parsed, out_path):
    doc = Document()
    for section_title, section_data in parsed.items():
        h = doc.add_heading(section_title.replace("_", " "), level=2)
        h.runs[0].font.color.rgb = RGBColor(0x1a, 0x3c, 0x5e)
        rows = _flat_lines(section_data)
        for label, value in rows:
            if label == "" and value == "":
                doc.add_paragraph("")
            elif value:
                p = doc.add_paragraph()
                run = p.add_run(f"{label.strip()}: ")
                run.bold = True
                p.add_run(value)
            elif label:
                p = doc.add_paragraph()
                p.add_run(label.strip()).bold = True
    doc.save(out_path)


# ── Batch processing ──────────────────────────────────────────────────────────
if __name__ == "__main__":
    RESUME_FOLDER = r"D:\Ktas Project\ATS\ATS Email Parser\Resume"
    EXTENSIONS    = {".pdf", ".docx", ".doc"}

    results = {}
    for fname in os.listdir(RESUME_FOLDER):
        if os.path.splitext(fname)[1].lower() not in EXTENSIONS:
            continue
        fpath = os.path.join(RESUME_FOLDER, fname)
        key   = os.path.splitext(fname)[0]
        print(f"Parsing: {fname}")
        try:
            parsed = parse_resume(fpath)
            results[key] = parsed

            # Export PDF
            export_to_pdf(parsed, os.path.join(RESUME_FOLDER, f"{key}_parsed.pdf"))
            # Export Word
            export_to_word(parsed, os.path.join(RESUME_FOLDER, f"{key}_parsed.docx"))
            # Export separate JSON per resume
            json_path = os.path.join(RESUME_FOLDER, f"{key}_parsed.json")
            with open(json_path, "w", encoding="utf-8") as jf:
                json.dump(parsed, jf, indent=4, ensure_ascii=False)

        except Exception as e:
            print(f"  ERROR: {e}")
            results[key] = {"error": str(e)}

    out_path = os.path.join(RESUME_FOLDER, "parsed_resumes.json")
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(results, f, indent=4, ensure_ascii=False)

    print(f"\nDone. {len(results)} resumes saved to: {out_path}")
